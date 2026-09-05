# -*- coding: utf-8 -*-
"""从 睡前故事/ 的 Markdown 源文件生成《睡前故事·亲子朗读手册》docx。
用法: python3 _build_handbook.py
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = Path(__file__).resolve().parent / "睡前故事·亲子朗读手册.docx"

PARTS = [
    ("第一辑 · 寓言", ROOT / "睡前故事" / "寓言"),
    ("第二辑 · 神话", ROOT / "睡前故事" / "神话"),
    ("第三辑 · 童话", ROOT / "睡前故事" / "童话"),
]

HEI, KAI, SONG = "黑体", "楷体", "宋体"
GRAY = RGBColor(0x59, 0x59, 0x59)


def set_east_asian(run, name):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def para(doc, text, font=SONG, size=12, bold=False, color=None,
         align=None, indent_chars=2, spacing=1.3, before=0, after=0,
         page_break_before=False):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.page_break_before = page_break_before
    if indent_chars:
        fmt.first_line_indent = Pt(size * indent_chars)
    if align is not None:
        fmt.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    set_east_asian(run, font)
    return p


def add_footer_page_number(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_east_asian(run, SONG)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* arabic \\* MERGEFORMAT")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    for f_run in p.runs:
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = GRAY


def parse_story(path):
    """返回 (标题, 信息行, 正文段列表, 晚安小结)"""
    text = path.read_text(encoding="utf-8")
    title, info, body, goodnight = "", "", [], ""
    in_body = False
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.startswith("# "):
            title = line[2:].strip()
        elif line.startswith(">"):
            content = line.lstrip("> ").strip()
            if content:
                info = (info + "　" + content).strip()
        elif line.startswith("**晚安小结**"):
            goodnight = line.replace("**晚安小结**：", "").replace("**晚安小结**:", "").strip()
            in_body = False
        elif line == "---":
            in_body = False
        elif line.strip():
            in_body = True
            body.append(line.strip())
    return title, info, body, goodnight


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.54)
        section.left_margin = section.right_margin = Cm(2.8)
    add_footer_page_number(doc)

    # 标题页
    para(doc, "", after=40)
    para(doc, "睡前故事", HEI, 30, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         indent_chars=0, after=6)
    para(doc, "亲 子 朗 读 手 册", HEI, 20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         indent_chars=0, after=18)
    para(doc, "寓言 · 神话 · 童话　共 30 篇", SONG, 13,
         align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, after=48, color=GRAY)
    for tip in [
        "每篇朗读约 3～5 分钟，正好在孩子入睡前讲完。",
        "正文用楷体大字排版，方便大人照读、孩子跟读。",
        "每篇末尾的「晚安小结」是给家长的话：读完轻声收尾，说一句晚安。",
        "推荐连播顺序、适读年龄与出处明细，见故事库《睡前故事索引》。",
    ]:
        para(doc, "· " + tip, SONG, 11.5, align=WD_ALIGN_PARAGRAPH.LEFT,
             indent_chars=0, after=6)

    story_no = 0
    for part_title, part_dir in PARTS:
        # 分辑扉页
        para(doc, "", after=100, page_break_before=True)
        para(doc, part_title, HEI, 22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             indent_chars=0)
        para(doc, "", after=0)

        for path in sorted(part_dir.glob("[0-9]*.md")):
            story_no += 1
            title, info, body, goodnight = parse_story(path)
            para(doc, f"{story_no:02d}　{title}", HEI, 16, bold=True,
                 indent_chars=0, after=4, page_break_before=True)
            para(doc, info, SONG, 10, color=GRAY, indent_chars=0, after=10)
            for seg in body:
                para(doc, seg, KAI, 12.5)
            if goodnight:
                p = para(doc, "晚安小结：", SONG, 11.5, bold=True,
                         indent_chars=0, before=10, after=0, color=GRAY)
                para(doc, goodnight, SONG, 11.5, indent_chars=0, color=GRAY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK 已生成: {OUT}（共 {story_no} 篇）")


if __name__ == "__main__":
    build()
