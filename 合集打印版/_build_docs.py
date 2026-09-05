# -*- coding: utf-8 -*-
"""从分类 Markdown 生成整册朗读手册 docx（通用版）。
用法:
  python3 _build_docs.py 成语故事     -> 成语故事·共读手册.docx
  python3 _build_docs.py 电台故事     -> 电台故事·深夜电台手册.docx
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
HEI, KAI, SONG = "黑体", "楷体", "宋体"
GRAY = RGBColor(0x59, 0x59, 0x59)

CONFIG = {
    "成语故事": {
        "subtitle": "成语背后的经典小故事",
        "tips": [
            "每篇开头注明成语含义与典籍出处，讲完可以让孩子复述成语。",
            "结尾「小启示」是给家长的话：照着和孩子聊两句。",
            "推荐组合与适读年龄明细，见故事库《成语故事索引》。",
        ],
        "glob": "[0-9]*.md",
    },
    "电台故事": {
        "subtitle": "深夜电台文本",
        "tips": [
            "每期按「开播的话 → 今晚的故事 → 收音前的话」三段式排版。",
            "适合深夜 FM 式朗读：语速放慢，段落间留呼吸，可配轻钢琴或白噪音。",
            "按情绪对照取用、连播建议，见故事库《电台故事索引》。",
        ],
        "glob": "第*.md",
    },
    "励志故事": {
        "subtitle": "名人轶事与励志小故事",
        "tips": [
            "涵盖中外人物与励志寓言，结尾「小启示」可读后与孩子共勉。",
            "李白、匡衡、达·芬奇、贝多芬、屠呦呦、刘伟……坚持的人各不相同。",
            "组合建议与适读年龄明细，见故事库《励志故事索引》。",
        ],
        "glob": "[0-9]*.md",
    },
    "科学故事": {
        "subtitle": "科学发现与发明背后的趣味故事",
        "tips": [
            "每篇注明人物、发现/发明与知识点，讲完可带孩子做延伸小科普。",
            "与励志故事互补：那边讲精神，这边讲科学是怎么发生的。",
            "组合建议与共读提示（万户篇结局较重），见故事库《科学趣味故事索引》。",
        ],
        "glob": "[0-9]*.md",
    },
}


def set_east_asian(run, name):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def para(doc, text, font=SONG, size=12, bold=False, color=None, align=None,
         indent_chars=0, spacing=1.3, before=0, after=0, page_break_before=False):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.page_break_before = page_break_before
    if indent_chars:
        fmt.first_line_indent = Pt(size * indent_chars)
    if align is not None:
        fmt.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    set_east_asian(run, font)
    return p


def add_footer_page_number(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* arabic \\* MERGEFORMAT")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    for f_run in p.runs:
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = GRAY


def parse_file(path: Path):
    """返回 (标题, 信息行, 正文段列表, 结尾标签, 结尾文本)"""
    text = path.read_text(encoding="utf-8").strip()
    lines = text.splitlines()

    title, title_seen = "", False
    meta, body = [], []
    ending_label, ending = "", ""

    for ln in lines:
        s = ln.rstrip()
        if not title_seen and s.startswith("# "):
            title = s[2:].strip()
            title_seen = True
            continue
        if not meta and s.startswith(">") and not body:
            meta.append(s.lstrip("> ").strip())
            continue
        m = re.match(r"^\*\*(晚安小结|小启示)\*\*[：:]\s*(.*)$", s)
        if m:
            ending_label, ending = m.group(1), m.group(2).strip()
            continue
        if s == "---":
            continue
        if s.strip():
            body.append(s)

    return title, "　".join(x for x in meta if x), body, ending_label, ending


def render_body(doc, segs, base_size=12):
    """正文渲染：## 小标题 / > 引用 / 普通段落"""
    buf = []

    def flush():
        if buf:
            para(doc, " ".join(buf), KAI, base_size, indent_chars=2)
            buf.clear()

    for seg in segs:
        s = seg.strip()
        if s.startswith("##"):
            flush()
            para(doc, s.lstrip("# ").strip(), HEI, 13.5, bold=True,
                 before=8, after=2)
        elif s.startswith(">"):
            flush()
            para(doc, s.lstrip("> ").strip(), SONG, 10.5, color=GRAY,
                 after=4)
        else:
            buf.append(s)
    flush()


def main():
    cat = sys.argv[1] if len(sys.argv) > 1 else ""
    if cat not in CONFIG:
        print(f"用法: python3 {Path(__file__).name} [{'/'.join(CONFIG)}]")
        sys.exit(1)
    cfg = CONFIG[cat]

    files = sorted((ROOT / cat).glob(cfg["glob"]))
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.54)
        section.left_margin = section.right_margin = Cm(2.8)
    add_footer_page_number(doc)

    # 标题页
    para(doc, "", after=40)
    para(doc, cat, HEI, 30, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    unit = "期" if cat == "电台故事" else "篇"
    sub = f"{cfg['subtitle']}　共 {len(files)} {unit}"
    para(doc, sub, HEI, 18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    para(doc, "故事集 · 朗读手册", SONG, 13, align=WD_ALIGN_PARAGRAPH.CENTER,
         after=48, color=GRAY)
    for tip in cfg["tips"]:
        para(doc, "· " + tip, SONG, 11.5, after=6)

    for i, f in enumerate(files):
        title, meta, body, end_label, ending = parse_file(f)
        para(doc, f"{i + 1:02d}　{title}", HEI, 16, bold=True,
             after=4, page_break_before=True)
        if meta:
            para(doc, meta, SONG, 10, color=GRAY, after=10)
        render_body(doc, body)
        if ending:
            para(doc, f"{end_label}：", SONG, 11.5, bold=True,
                 before=10, after=0, color=GRAY)
            para(doc, ending, SONG, 11.5, color=GRAY)

    out = Path(__file__).resolve().parent / f"{cat}·朗读手册.docx"
    doc.save(out)
    print(f"OK 已生成: {out}（共 {len(files)} 篇）")


if __name__ == "__main__":
    main()
